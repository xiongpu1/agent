"""
Word文档内容提取模块
按照文档顺序提取Word的全部内容（文本、表格、图片）
输出格式为json，由于是python-docx库提取的，是按照文档顺序排列的，但是没有形成段落内容，需要自己进行处理
图片内容是base64编码的，需要自己进行处理
"""

from docx import Document
import zipfile, os, base64, json
from typing import Dict, List, Optional, Any

def extract_content_from_docx(file_path: str) -> Optional[str]:
    """
    使用python-docx提取Word文档的完整内容，按文档顺序排列
    
    Args:
        file_path (str): docx文件路径
    
    Returns:
        str: JSON格式的字符串，包含document_info、ordered_content、images（含base64编码）
    """
    if not os.path.exists(file_path):
        print(f"错误: 文件 {file_path} 不存在")
        return None
    
    try:
        # 打开docx文件
        doc = Document(file_path)
        
        print(f"正在提取文件: {file_path}")
        
        # 提取文档信息
        doc_info = {
            "filename": os.path.basename(file_path),
            "paragraphs_count": len(doc.paragraphs),
            "tables_count": len(doc.tables),
            "sections_count": len(doc.sections)
        }
        
        # 按顺序提取所有内容元素
        ordered_content = []
        
        # 遍历文档的所有元素（段落和表格）
        for element in doc.element.body:
            if element.tag.endswith('p'):  # 段落
                # 找到对应的段落对象
                for para in doc.paragraphs:
                    if para._element == element:
                        if para.text.strip():  # 只保存非空段落
                            para_info = {
                                "type": "paragraph",
                                "index": len(ordered_content) + 1,
                                "text": para.text.strip(),
                                "style": para.style.name if para.style else "Normal",
                                "alignment": str(para.alignment) if para.alignment else "None"
                            }
                            ordered_content.append(para_info)
                        break
            
            elif element.tag.endswith('tbl'):  # 表格
                # 找到对应的表格对象
                for table in doc.tables:
                    if table._element == element:
                        table_data = {
                            "type": "table",
                            "index": len(ordered_content) + 1,
                            "rows": len(table.rows),
                            "columns": len(table.columns),
                            "cells": []
                        }
                        
                        for row_idx, row in enumerate(table.rows):
                            row_data = []
                            for cell_idx, cell in enumerate(row.cells):
                                cell_info = {
                                    "row": row_idx + 1,
                                    "column": cell_idx + 1,
                                    "text": cell.text.strip()
                                }
                                row_data.append(cell_info)
                            table_data["cells"].append(row_data)
                        
                        ordered_content.append(table_data)
                        break
        
        # 提取图片信息
        images = extract_images_from_docx(file_path)
        
        # 组合所有内容
        content = {
            "document_info": doc_info,
            "ordered_content": ordered_content,
            "images": images
        }
        
        print(f"提取完成:")
        print(f"  总元素数: {len(ordered_content)}")
        print(f"  图片数: {len(images)}")
        
        # 转换为JSON字符串
        json_content = json.dumps(content, ensure_ascii=False, indent=2)
        return json_content
        
    except Exception as e:
        print(f"提取文件时发生错误: {e}")
        return None


def extract_images_from_docx(file_path: str) -> List[Dict[str, Any]]:
    """
    从docx文件中提取图片信息，包含base64编码
    
    Args:
        file_path (str): docx文件路径
    
    Returns:
        list: 图片信息列表，包含base64编码
    """
    images = []
    
    try:
        # 打开docx文件作为zip文件
        with zipfile.ZipFile(file_path, 'r') as docx_zip:
            # 查找图片文件
            image_files = [f for f in docx_zip.namelist() if f.startswith('word/media/') and f != 'word/media/']
            
            for img_file in image_files:
                try:
                    # 读取图片数据
                    img_data = docx_zip.read(img_file)
                    
                    # 获取文件扩展名
                    _, ext = os.path.splitext(img_file)
                    if not ext:
                        ext = '.jpeg'  # 默认扩展名

                    img_base64 = base64.b64encode(img_data).decode('utf-8')
                    
                    # 创建图片信息字典
                    image_info = {
                        "filename": os.path.basename(img_file),
                        "path": img_file,
                        "size": len(img_data),
                        "extension": ext,
                        "size_mb": round(len(img_data) / (1024 * 1024), 2),
                        "base64": img_base64
                    }
                    images.append(image_info)
                    
                except Exception as e:
                    print(f"读取图片 {img_file} 时出错: {e}")
        
        return images
        
    except Exception as e:
        print(f"提取图片时出错: {e}")
        return []

