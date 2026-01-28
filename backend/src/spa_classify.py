from __future__ import annotations

import argparse
import base64
import io
import json
import os
import re
import time
import urllib.error
import urllib.request
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Set, Tuple

try:
    from dotenv import load_dotenv
except Exception:  # pragma: no cover
    load_dotenv = None  # type: ignore

try:
    # Run with: python -m src.spa_classify (cwd=backend)
    from src.kb_chat.llm_client import chat_json  # type: ignore
except Exception:  # pragma: no cover
    try:
        # Run with: python spa_classify.py (cwd=backend/src)
        from kb_chat.llm_client import chat_json  # type: ignore
    except Exception:  # pragma: no cover
        chat_json = None  # type: ignore

from openai import OpenAI


@dataclass
class FileItem:
    abs_path: Path
    rel_path: str
    ext: str
    size: int
    mtime: float


_TEXT_EXTS = {".md", ".markdown", ".txt", ".json"}
_EXCEL_EXTS = {".xlsx", ".xls"}
_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tif", ".tiff", ".webp", ".jfif", ".pjpeg", ".pjp"}
_PDF_EXTS = {".pdf"}
_DOCX_EXTS = {".docx"}


def _iter_files(root: Path) -> Iterable[FileItem]:
    root = root.resolve()
    for p in root.rglob("*"):
        if not p.is_file():
            continue
        try:
            st = p.stat()
        except Exception:
            continue
        rel = str(p.relative_to(root))
        ext = p.suffix.lower()
        yield FileItem(abs_path=p, rel_path=rel, ext=ext, size=int(st.st_size), mtime=float(st.st_mtime))


def _read_text_snippet(path: Path, max_chars: int) -> str:
    if max_chars <= 0:
        return ""

    encodings = ["utf-8", "utf-8-sig", "gb18030"]
    data: Optional[str] = None
    for enc in encodings:
        try:
            data = path.read_text(encoding=enc, errors="strict")
            break
        except Exception:
            continue

    if data is None:
        try:
            data = path.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            return ""

    data = data.replace("\u0000", " ")
    return data[:max_chars]


def _read_excel_snippet(path: Path, max_chars: int, nrows: int = 20) -> str:
    try:
        import pandas as pd
    except Exception:
        return ""

    try:
        df = pd.read_excel(path, nrows=nrows)
    except Exception:
        return ""

    try:
        md = df.to_markdown(index=False)
    except Exception:
        md = str(df.head(nrows))

    md = md.strip()
    return md[:max_chars]


def _read_docx_snippet(path: Path, max_chars: int) -> str:
    try:
        from docx import Document  # type: ignore
    except Exception:
        return ""

    try:
        doc = Document(str(path))
    except Exception:
        return ""

    parts: List[str] = []
    try:
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)
    except Exception:
        return ""

    try:
        for table in doc.tables:
            for row in table.rows:
                row_texts: List[str] = []
                for cell in row.cells:
                    ct = (cell.text or "").strip()
                    if ct:
                        row_texts.append(ct)
                if row_texts:
                    parts.append(" | ".join(row_texts))
    except Exception:
        pass

    text = "\n".join(parts).strip()
    return text[:max_chars]


def _extract_first_docx_image_data_url(path: Path, max_bytes: int) -> Optional[str]:
    try:
        from docx import Document  # type: ignore
    except Exception:
        return None

    try:
        doc = Document(str(path))
    except Exception:
        return None

    try:
        rels = getattr(doc.part, "related_parts", {})
        for _, part in rels.items():
            ct = str(getattr(part, "content_type", "") or "")
            if not ct.startswith("image/"):
                continue
            blob = getattr(part, "blob", None)
            if not blob:
                continue
            b = bytes(blob)
            if max_bytes > 0 and len(b) > max_bytes:
                continue
            b64 = base64.b64encode(b).decode("utf-8")
            return f"data:{ct};base64,{b64}"
    except Exception:
        return None

    return None


_MD_INLINE_DATA_IMAGE_RE = re.compile(r"!\[[^\]]*\]\(data:image/[^)]+\)")


def _sanitize_markdown_for_text_llm(md: str) -> str:
    t = (md or "").strip()
    if not t:
        return ""
    t = _MD_INLINE_DATA_IMAGE_RE.sub("[image]", t)
    return t


def _try_extract_text_with_fitz(pdf_path: Path) -> Optional[str]:
    try:
        import fitz  # type: ignore
    except Exception:
        return None

    try:
        doc = fitz.open(str(pdf_path))
        parts: List[str] = []
        for page in doc:
            parts.append(page.get_text("text") or "")
        out = "\n".join(parts).strip()
        return out or None
    except Exception:
        return None


def _try_render_pdf_first_page_data_url(pdf_path: Path, max_bytes: int) -> Optional[str]:
    try:
        import fitz  # type: ignore
    except Exception:
        return None

    try:
        doc = fitz.open(str(pdf_path))
        if doc.page_count <= 0:
            return None
        page = doc.load_page(0)
        pix = page.get_pixmap(dpi=150)
        b = pix.tobytes("png")
        if max_bytes > 0 and len(b) > max_bytes:
            return None
        b64 = base64.b64encode(b).decode("utf-8")
        return f"data:image/png;base64,{b64}"
    except Exception:
        return None


def _try_extract_text_with_pypdf(pdf_path: Path) -> Optional[str]:
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception:
        return None

    try:
        reader = PdfReader(str(pdf_path))
        parts: List[str] = []
        for page in reader.pages:
            parts.append(page.extract_text() or "")
        out = "\n".join(parts).strip()
        return out or None
    except Exception:
        return None


def _try_extract_text_with_pypdf2(pdf_path: Path) -> Optional[str]:
    try:
        from PyPDF2 import PdfReader  # type: ignore
    except Exception:
        return None

    try:
        reader = PdfReader(str(pdf_path))
        parts: List[str] = []
        for page in reader.pages:
            parts.append(page.extract_text() or "")
        out = "\n".join(parts).strip()
        return out or None
    except Exception:
        return None


def _read_pdf_excerpt_and_image(
    path: Path,
    max_chars: int,
    max_pages: int,
    lang: str,
    max_image_bytes: int,
) -> Tuple[str, Optional[str]]:
    parse_doc = None
    try:
        from src.data_pdf import parse_doc as _parse_doc  # type: ignore

        parse_doc = _parse_doc
    except Exception:
        try:
            from data_pdf import parse_doc as _parse_doc  # type: ignore

            parse_doc = _parse_doc
        except Exception:
            parse_doc = None

    try:
        mp = int(max_pages)
    except Exception:
        mp = 3
    if mp <= 0:
        mp = 3

    md_raw_s = ""
    first_image: Optional[str] = None
    end_page_id = mp - 1
    if parse_doc is not None:
        try:
            md_raw = parse_doc(path, lang=lang, backend="pipeline", method="auto", start_page_id=0, end_page_id=end_page_id)
            md_raw_s = "" if not md_raw else str(md_raw)
            first_image = _extract_first_data_image_url(md_raw_s)
        except Exception:
            md_raw_s = ""
            first_image = None

    md = _sanitize_markdown_for_text_llm(md_raw_s)
    md = md.strip()
    if md:
        return (md[:max_chars], first_image)

    # Fallback: try to extract text from the PDF text layer.
    t = _try_extract_text_with_fitz(path)
    if not t:
        t = _try_extract_text_with_pypdf(path)
    if not t:
        t = _try_extract_text_with_pypdf2(path)
    if not t:
        rendered = _try_render_pdf_first_page_data_url(path, max_bytes=int(max_image_bytes))
        return ("", first_image or rendered)

    return (str(t)[:max_chars], first_image)


_MD_ANY_DATA_IMAGE_URL_RE = re.compile(r"\(data:image/[^)]+\)")


def _extract_first_data_image_url(md: str) -> Optional[str]:
    t = md or ""
    m = _MD_ANY_DATA_IMAGE_URL_RE.search(t)
    if not m:
        return None
    frag = m.group(0)
    if frag.startswith("(") and frag.endswith(")"):
        frag = frag[1:-1]
    frag = frag.strip()
    return frag or None


def _to_data_url_for_image(path: Path, max_bytes: int) -> Optional[str]:
    """
    将图片转换为 data URL，如果超过大小限制则自动压缩
    
    Args:
        path: 图片文件路径
        max_bytes: 最大字节数（0 表示不限制）
        
    Returns:
        data URL 字符串，失败返回 None
    """
    try:
        from PIL import Image
        import io
        
        # 读取原始图片
        img = Image.open(path)
        
        # 转换为 RGB（如果是 RGBA 或其他模式）
        if img.mode in ('RGBA', 'LA', 'P'):
            # 创建白色背景
            background = Image.new('RGB', img.size, (255, 255, 255))
            if img.mode == 'P':
                img = img.convert('RGBA')
            background.paste(img, mask=img.split()[-1] if img.mode in ('RGBA', 'LA') else None)
            img = background
        elif img.mode != 'RGB':
            img = img.convert('RGB')
        
        # 获取原始大小
        original_size = path.stat().st_size
        
        # 如果不限制大小，或者原始文件已经小于限制，直接返回
        if max_bytes <= 0 or original_size <= max_bytes:
            b = path.read_bytes()
            ext = path.suffix.lower()
            mime = "image/jpeg"
            if ext == ".png":
                mime = "image/png"
            elif ext == ".webp":
                mime = "image/webp"
            elif ext == ".gif":
                mime = "image/gif"
            b64 = base64.b64encode(b).decode("utf-8")
            return f"data:{mime};base64,{b64}"
        
        # 需要压缩：逐步降低质量直到满足大小要求
        quality = 85
        scale = 1.0
        
        while quality >= 20 or scale > 0.3:
            # 调整尺寸
            if scale < 1.0:
                new_size = (int(img.width * scale), int(img.height * scale))
                resized_img = img.resize(new_size, Image.Resampling.LANCZOS)
            else:
                resized_img = img
            
            # 压缩为 JPEG
            buffer = io.BytesIO()
            resized_img.save(buffer, format='JPEG', quality=quality, optimize=True)
            compressed_bytes = buffer.getvalue()
            
            # 检查大小
            if len(compressed_bytes) <= max_bytes:
                # 成功压缩到目标大小
                b64 = base64.b64encode(compressed_bytes).decode("utf-8")
                print(f"    图片已压缩: {original_size / 1024 / 1024:.1f}MB → {len(compressed_bytes) / 1024 / 1024:.1f}MB (质量={quality}, 缩放={scale:.0%})")
                return f"data:image/jpeg;base64,{b64}"
            
            # 继续降低质量或缩小尺寸
            if quality > 20:
                quality -= 10
            else:
                scale -= 0.1
        
        # 如果还是太大，返回 None
        print(f"    ⚠️ 图片无法压缩到目标大小: {original_size / 1024 / 1024:.1f}MB")
        return None
        
    except Exception as e:
        print(f"    ⚠️ 图片处理失败: {e}")
        return None


def _kind_for_ext(ext: str) -> str:
    if ext in _TEXT_EXTS:
        return "text"
    if ext in _EXCEL_EXTS:
        return "excel"
    if ext in _IMAGE_EXTS:
        return "image"
    if ext in _PDF_EXTS:
        return "pdf"
    if ext in _DOCX_EXTS:
        return "docx"
    return "other"


_JSON_RE = re.compile(r"\{.*\}", re.DOTALL)


def _extract_json_obj(text: str) -> Optional[Dict[str, Any]]:
    t = (text or "").strip()
    if not t:
        return None

    m = _JSON_RE.search(t)
    if not m:
        return None

    snippet = m.group(0)
    try:
        obj = json.loads(snippet)
    except Exception:
        return None

    if not isinstance(obj, dict):
        return None

    return obj


def _resolve_vision_model(explicit: Optional[str]) -> str:
    if (explicit or "").strip():
        return str(explicit).strip()
    m = (os.getenv("POSTER_VISION_MODEL", "") or "").strip()
    if m:
        return m
    m = (os.getenv("DEFAULT_VISION_MODEL", "") or "").strip()
    if m:
        return m
    return "dashscope/qwen-vl-max"


def _normalize_dashscope_model_name(model: str) -> str:
    m = (model or "").strip()
    if m.startswith("dashscope/"):
        return m.split("/", 1)[1]
    return m


_OPENAI_CLIENT: Optional[OpenAI] = None


def _get_openai_client() -> OpenAI:
    global _OPENAI_CLIENT
    if _OPENAI_CLIENT is not None:
        return _OPENAI_CLIENT

    api_key = os.getenv("DASHSCOPE_API_KEY")
    base_url = os.getenv("DASHSCOPE_BASE_URL")
    if not (base_url or "").strip():
        base_url = "https://dashscope.aliyuncs.com/compatible-mode/v1"

    _OPENAI_CLIENT = OpenAI(api_key=api_key, base_url=base_url)
    return _OPENAI_CLIENT


def _is_ollama_base_url(base_url: str) -> bool:
    u = (base_url or "").strip().lower()
    if not u:
        return False
    # Heuristic: default Ollama port or localhost usage.
    return ("11434" in u) or ("localhost" in u) or ("127.0.0.1" in u)


def _ollama_api_chat_url() -> str:
    base_url = (os.getenv("DASHSCOPE_BASE_URL") or "").strip()
    if not base_url:
        base_url = "http://127.0.0.1:11434/v1"
    # Convert OpenAI-compatible base url (/v1) to Ollama native /api/chat.
    u = base_url.rstrip("/")
    if u.endswith("/v1"):
        u = u[: -len("/v1")]
    return u.rstrip("/") + "/api/chat"


def _data_url_to_b64_payload(data_url: str) -> Optional[str]:
    t = (data_url or "").strip()
    if not t:
        return None
    if t.startswith("data:"):
        # data:image/jpeg;base64,.....
        if "," not in t:
            return None
        return t.split(",", 1)[1].strip() or None
    # If caller passed a raw base64 string already.
    return t


def _ollama_reencode_image_to_jpeg_b64(image_b64: str, max_bytes: int) -> Optional[str]:
    try:
        from PIL import Image, ImageFile  # type: ignore
    except Exception:
        return image_b64

    try:
        raw = base64.b64decode(image_b64)
    except Exception:
        return None

    try:
        ImageFile.LOAD_TRUNCATED_IMAGES = True
        img = Image.open(io.BytesIO(raw))
        img.load()
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        elif img.mode == "L":
            img = img.convert("RGB")

        out = io.BytesIO()
        # Force baseline JPEG (progressive=False) to maximize decoder compatibility.
        img.save(out, format="JPEG", quality=92, optimize=True, progressive=False)
        b = out.getvalue()
        if max_bytes > 0 and len(b) > max_bytes:
            return None
        return base64.b64encode(b).decode("utf-8")
    except Exception:
        pass

    try:
        import cv2  # type: ignore
        import numpy as np  # type: ignore
    except Exception:
        return None

    try:
        arr = np.frombuffer(raw, dtype=np.uint8)
        img2 = cv2.imdecode(arr, cv2.IMREAD_COLOR)
        if img2 is None:
            return None
        ok, enc = cv2.imencode(".jpg", img2, [int(cv2.IMWRITE_JPEG_QUALITY), 92])
        if not ok:
            return None
        b2 = enc.tobytes()
        if max_bytes > 0 and len(b2) > max_bytes:
            return None
        return base64.b64encode(b2).decode("utf-8")
    except Exception:
        return None


def _ollama_chat_json_with_image(model: str, system: str, user_text: str, image_b64: str, timeout_s: int) -> Dict[str, Any]:
    url = _ollama_api_chat_url()
    payload = {
        "model": model,
        "stream": False,
        "format": "json",
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user_text, "images": [image_b64]},
        ],
    }
    req = urllib.request.Request(
        url,
        data=json.dumps(payload).encode("utf-8"),
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=timeout_s) as resp:
        raw = resp.read().decode("utf-8", errors="replace")
    try:
        return json.loads(raw)
    except Exception:
        return {"_raw": raw}


def _ollama_chat_json(messages: List[Dict[str, Any]], model: Optional[str], timeout_s: int) -> Dict[str, Any]:
    url = _ollama_api_chat_url()
    m = (model or os.getenv("DEFAULT_LLM_MODEL") or "").strip() or "qwen3:32b"
    # In case caller passed dashscope/xxx
    m = _normalize_dashscope_model_name(m)
    payload = {
        "model": m,
        "stream": False,
        "format": "json",
        "messages": messages,
    }
    req = urllib.request.Request(
        url,
        data=json.dumps(payload).encode("utf-8"),
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=timeout_s) as resp:
        raw = resp.read().decode("utf-8", errors="replace")
    try:
        obj = json.loads(raw)
    except Exception:
        # Preserve raw so caller can debug.
        return {"choices": [{"message": {"content": raw}}]}

    content = str((obj.get("message") or {}).get("content") or "").strip()
    if not content and "error" in obj:
        content = json.dumps({"error": obj.get("error")}, ensure_ascii=False)
    return {"choices": [{"message": {"content": content}}]}


def _chat_json_fallback_openai(messages: List[Dict[str, Any]], model: Optional[str]) -> Dict[str, Any]:
    client = _get_openai_client()
    m = (model or os.getenv("DEFAULT_LLM_MODEL") or "").strip() or "qwen3-max"
    m = _normalize_dashscope_model_name(m)
    resp = client.chat.completions.create(model=m, messages=messages)
    content = (resp.choices[0].message.content or "").strip()
    return {
        "choices": [{"message": {"content": content}}],
        "model": getattr(resp, "model", None),
        "usage": getattr(resp, "usage", None),
        "id": getattr(resp, "id", None),
    }


def _chat_json(messages: List[Dict[str, Any]], model: Optional[str]) -> Dict[str, Any]:
    base_url = (os.getenv("DASHSCOPE_BASE_URL") or "").strip()
    if _is_ollama_base_url(base_url):
        timeout_s = 300
        try:
            timeout_s = int(os.getenv("KBCHAT_LLM_TIMEOUT_SECONDS", "300"))
        except Exception:
            timeout_s = 300
        try:
            return _ollama_chat_json(messages=messages, model=model, timeout_s=timeout_s)
        except Exception:
            # Fall back to OpenAI-compatible path.
            pass

    if callable(chat_json):
        try:
            return chat_json(messages=messages, model=model)  # type: ignore
        except Exception:
            pass
    return _chat_json_fallback_openai(messages=messages, model=model)


def _resolve_fallback_vision_model(primary: str) -> Optional[str]:
    m = (os.getenv("FALLBACK_VISION_MODEL", "") or "").strip()
    if m and m != primary:
        return m
    if primary != "dashscope/qwen-vl-max":
        return "dashscope/qwen-vl-max"
    return None


def _response_suggests_missing_image(content: str) -> bool:
    t = (content or "").strip()
    if not t:
        return True
    # Common pattern when the image part is ignored by the server.
    if "请提供图片" in t or "无法看到图片" in t or "看不到图片" in t:
        return True
    return False


def _build_dashscope_kwargs(model: Optional[str]) -> Dict[str, Any]:
    m = _resolve_vision_model(model)
    timeout_s = 60
    try:
        timeout_s = int(os.getenv("KBCHAT_LLM_TIMEOUT_SECONDS", "60"))
    except Exception:
        timeout_s = 60

    kwargs: Dict[str, Any] = {
        "model": m,
        "temperature": 0.2,
        "timeout": timeout_s,
        "request_timeout": timeout_s,
        "api_base": "https://dashscope.aliyuncs.com/compatible-mode/v1",
    }
    api_key = os.getenv("DASHSCOPE_API_KEY")
    if api_key:
        kwargs["api_key"] = api_key
    return kwargs


def _format_meta_text(meta: Optional[Dict[str, Any]]) -> str:
    if not meta:
        return ""
    parts: List[str] = []
    rel_path = str(meta.get("rel_path", "") or "").strip()
    file_name = str(meta.get("file_name", "") or "").strip()
    ext = str(meta.get("ext", "") or "").strip()
    size = meta.get("size")
    mtime = meta.get("mtime")

    if rel_path:
        parts.append(f"rel_path: {rel_path}")
    if file_name:
        parts.append(f"file_name: {file_name}")
    if ext:
        parts.append(f"ext: {ext}")
    if size is not None:
        parts.append(f"size: {size}")
    if mtime is not None:
        parts.append(f"mtime: {mtime}")

    if not parts:
        return ""
    return "文件元信息(可用于理解用途/编号):\n" + "\n".join(parts) + "\n"


def _vision_capsule(data_url: str, model: Optional[str], meta: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    system = (
        "你是一个企业资料库的图片内容分析助手。"
        "请基于图片内容，并结合提供的文件元信息(如果有)生成结构化结果。输出必须是严格 JSON："
        "{\"summary\": string, \"keyphrases\": string[], \"confidence_read\": number}."
        "\n"
        "要求：\n"
        "- summary: 50-120 个中文字符，描述图片的核心内容/用途（例如产品外观、场景展示、配色、细节等）。\n"
        "- keyphrases: 3-8 个中文或英文短词（型号/系列/颜色/材质/部位/场景等）。\n"
        "- confidence_read: 0~1，表示你对图片内容理解的把握。\n"
        "- 允许参考文件名/路径中的编号等线索，但不要臆测不存在的信息。"
    )

    # Keep ordering consistent with DashScope OpenAI-compatible examples:
    # image_url first, then text.
    meta_text = _format_meta_text(meta)
    user = [
        {"type": "image_url", "image_url": {"url": data_url}},
        {"type": "text", "text": f"{meta_text}请按要求输出 JSON。"},
    ]

    resolved_model = _resolve_vision_model(model)
    base_url = (os.getenv("DASHSCOPE_BASE_URL") or "").strip()

    content = ""
    if _is_ollama_base_url(base_url):
        img_b64 = _data_url_to_b64_payload(data_url)
        if not img_b64:
            return {"summary": "视觉模型输入图片解析失败。", "keyphrases": [], "confidence_read": 0.0}

        try:
            max_bytes = int(os.getenv("OLLAMA_IMAGE_MAX_BYTES", "6000000"))
        except Exception:
            max_bytes = 6_000_000
        img_b64 = _ollama_reencode_image_to_jpeg_b64(img_b64, max_bytes=max_bytes)
        if not img_b64:
            return {"summary": "视觉模型输入图片重编码失败或图片过大。", "keyphrases": [], "confidence_read": 0.0}

        timeout_s = 300
        try:
            timeout_s = int(os.getenv("KBCHAT_LLM_TIMEOUT_SECONDS", "300"))
        except Exception:
            timeout_s = 300
        user_text = f"{meta_text}请按要求输出 JSON。"
        try:
            resp_obj = _ollama_chat_json_with_image(
                model=str(resolved_model),
                system=system,
                user_text=user_text,
                image_b64=img_b64,
                timeout_s=timeout_s,
            )
            content = str((resp_obj.get("message") or {}).get("content") or "").strip()
            if not content and "error" in resp_obj:
                return {"summary": f"视觉模型调用失败: {resp_obj.get('error')}", "keyphrases": [], "confidence_read": 0.0}
            if not content and "_raw" in resp_obj:
                return {"summary": f"视觉模型调用失败: {str(resp_obj.get('_raw'))[:200]}", "keyphrases": [], "confidence_read": 0.0}
        except urllib.error.HTTPError as e:
            try:
                err = e.read().decode("utf-8", errors="replace")
            except Exception:
                err = ""
            return {"summary": f"视觉模型调用失败: HTTPError({e.code}) {err[:200]}", "keyphrases": [], "confidence_read": 0.0}
        except Exception as e:
            return {"summary": f"视觉模型调用失败: {repr(e)}", "keyphrases": [], "confidence_read": 0.0}
    else:
        primary_model = _normalize_dashscope_model_name(resolved_model)
        client = _get_openai_client()

        try:
            resp = client.chat.completions.create(
                model=primary_model,
                messages=[
                    {"role": "system", "content": system},
                    {"role": "user", "content": user},
                ],
            )
        except Exception as e:
            return {"summary": f"视觉模型调用失败: {repr(e)}", "keyphrases": [], "confidence_read": 0.0}

        content = (resp.choices[0].message.content or "").strip()

    # If the server ignored the image payload, content often asks user to provide images.
    # We keep this function single-shot to avoid surprise extra API calls.
    if _response_suggests_missing_image(content):
        fb = _resolve_fallback_vision_model(_normalize_dashscope_model_name(resolved_model))
        if fb:
            return {
                "summary": "视觉模型未接收到图片（可能不支持 data: base64 图片输入）；建议改用支持的视觉模型或提供公网图片 URL。",
                "keyphrases": ["missing_image"],
                "confidence_read": 0.0,
                "_debug_primary_model": str(resolved_model),
                "_debug_fallback_suggestion": fb,
            }

    obj = _extract_json_obj(content) or {}
    summary = str(obj.get("summary", "")).strip()
    if summary:
        summary = summary.replace("\n", " ").strip()
        if len(summary) > 160:
            summary = summary[:160]

    keyphrases_raw = obj.get("keyphrases", [])
    keyphrases: List[str] = []
    if isinstance(keyphrases_raw, list):
        keyphrases = [str(x).strip() for x in keyphrases_raw if str(x).strip()]
    else:
        keyphrases = [str(keyphrases_raw).strip()] if str(keyphrases_raw).strip() else []

    conf_raw = obj.get("confidence_read", 0.0)
    try:
        confidence = float(conf_raw)
    except Exception:
        confidence = 0.0
    confidence = max(0.0, min(confidence, 1.0))

    if not summary:
        summary = "图片内容摘要生成失败。"

    return {"summary": summary, "keyphrases": keyphrases[:12], "confidence_read": confidence}


def _text_capsule(kind: str, excerpt: str, model: Optional[str], meta: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    system = (
        "你是一个企业资料库的文档/表格内容总结助手。"
        "请基于提供的内容摘录生成结构化结果。输出必须是严格 JSON："
        "{\"summary\": string, \"keyphrases\": string[], \"confidence_read\": number}."
        "\n"
        "要求：\n"
        "- summary: 50-120 个中文字符，描述文件核心内容/用途。\n"
        "- keyphrases: 3-8 个短词（字段名/主题/产品型号/关键指标等）。\n"
        "- confidence_read: 0~1，表示你对摘录足以概括全文的把握。\n"
        "- 可结合文件元信息(如果提供)中的编号/关键字辅助理解，但不要臆测不存在的信息。"
    )

    meta_text = _format_meta_text(meta)
    user = (
        f"{meta_text}"
        f"内容类型: {kind}\n"
        f"内容摘录(可能不完整):\n{excerpt}\n"
        "请只返回 JSON。"
    )

    resp = _chat_json(
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        model=model,
    )

    content = ""
    try:
        content = resp["choices"][0]["message"]["content"]
    except Exception:
        content = str(resp)

    obj = _extract_json_obj(content) or {}
    summary = str(obj.get("summary", "")).strip()
    if summary:
        summary = summary.replace("\n", " ").strip()
        if len(summary) > 160:
            summary = summary[:160]

    keyphrases_raw = obj.get("keyphrases", [])
    keyphrases: List[str] = []
    if isinstance(keyphrases_raw, list):
        keyphrases = [str(x).strip() for x in keyphrases_raw if str(x).strip()]
    else:
        keyphrases = [str(keyphrases_raw).strip()] if str(keyphrases_raw).strip() else []

    conf_raw = obj.get("confidence_read", 0.0)
    try:
        confidence = float(conf_raw)
    except Exception:
        confidence = 0.0
    confidence = max(0.0, min(confidence, 1.0))

    if not summary:
        summary = "内容摘要生成失败。"

    return {"summary": summary, "keyphrases": keyphrases[:12], "confidence_read": confidence}


def _llm_classify(rel_path: str, kind: str, ext: str, excerpt: str, model: Optional[str]) -> Dict[str, Any]:
    system = (
        "你是一个企业资料库的文件分类与摘要助手。"
        "对每个文件做内容级分类（不要依赖文件夹/目录名做判断）。"
        "你的输出必须是严格 JSON，对象字段固定如下："
        "{\"category\": string, \"summary\": string, \"confidence\": number, \"evidence\": string}."
        "\n"
        "要求：\n"
        "- category: 2-8 个中文字符，自发现类目，尽量稳定可复用（例如：产品图片、规格页、证书、生产入库表、安装说明、售后、营销素材、视频等）。\n"
        "- summary: 50-120 个中文字符，描述文件核心内容；不要输出项目符号列表。\n"
        "- confidence: 0~1 的小数。\n"
        "- evidence: 用 1-2 句话说明你为什么这么分（引用摘录中的关键短语/字段名即可）。\n"
        "- 如果信息不足，请给出最可能的 category，并在 evidence 里说明信息不足。"
    )

    user = (
        f"文件标识: {os.path.basename(rel_path)}\n"
        f"文件类型: {kind}\n"
        f"扩展名: {ext}\n"
        f"内容摘录(可能为空):\n{excerpt}\n"
        "请只返回 JSON，不要任何多余文字。"
    )

    resp = _chat_json(
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        model=model,
    )

    content = ""
    try:
        content = resp["choices"][0]["message"]["content"]
    except Exception:
        content = str(resp)

    obj = _extract_json_obj(content)
    if obj is None:
        return {
            "category": "未识别",
            "summary": "无法从模型返回中解析结构化结果。",
            "confidence": 0.0,
            "evidence": (content or "")[:200],
        }

    category = str(obj.get("category", "")).strip() or "未分类"
    summary = str(obj.get("summary", "")).strip()
    evidence = str(obj.get("evidence", "")).strip()

    if summary:
        summary = summary.replace("\n", " ").strip()
        if len(summary) > 160:
            summary = summary[:160]

    conf_raw = obj.get("confidence", 0.0)
    try:
        confidence = float(conf_raw)
    except Exception:
        confidence = 0.0
    confidence = max(0.0, min(confidence, 1.0))

    return {
        "category": category,
        "summary": summary,
        "confidence": confidence,
        "evidence": evidence,
    }


def _llm_classify_from_capsule(
    modality: str,
    summary: str,
    keyphrases: List[str],
    model: Optional[str],
    meta: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    # 导入分类体系
    try:
        from category_taxonomy import get_category_prompt, validate_category, get_closest_category
    except ImportError:
        try:
            from src.category_taxonomy import get_category_prompt, validate_category, get_closest_category
        except ImportError:
            # 如果导入失败，使用旧逻辑
            return _llm_classify_from_capsule_old(modality, summary, keyphrases, model, meta)
    
    category_prompt = get_category_prompt()
    
    system = (
        "你是一个企业资料库的文件分类助手。"
        "请基于提供的内容胶囊(capsule)，并结合文件元信息(如果提供)进行分类。"
        "输出必须是严格 JSON：{\"category_l1\": string, \"category_l2\": string, \"confidence\": number, \"evidence\": string}."
        "\n"
        f"{category_prompt}"
        "\n"
        "要求：\n"
        "- category_l1: 必须从上述预定义的一级类目中选择一个。\n"
        "- category_l2: 必须从对应一级类目下的二级类目中选择一个。\n"
        "- confidence: 0~1，表示分类的置信度。\n"
        "- evidence: 引用 summary/keyphrases 或 文件元信息中的关键字来说明理由。\n"
        "- 如果不确定，选择最接近的类别，并在 evidence 中说明不确定的原因。"
    )

    meta_text = _format_meta_text(meta)
    user = (
        f"{meta_text}"
        f"模态: {modality}\n"
        f"summary: {summary}\n"
        f"keyphrases: {json.dumps(keyphrases, ensure_ascii=False)}\n"
        "请只返回 JSON。"
    )

    resp = _chat_json(
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        model=model,
    )

    content = ""
    try:
        content = resp["choices"][0]["message"]["content"]
    except Exception:
        content = str(resp)

    obj = _extract_json_obj(content)
    if obj is None:
        return {
            "category_l1": "内部管理",
            "category_l2": "其他文档",
            "category": "内部管理/其他文档",
            "confidence": 0.0,
            "evidence": (content or "")[:200],
        }

    category_l1 = str(obj.get("category_l1", "")).strip() or "内部管理"
    category_l2 = str(obj.get("category_l2", "")).strip() or "其他文档"
    
    # 验证类别是否在预定义体系中
    if not validate_category(category_l1, category_l2):
        # 找到最接近的类别
        category_l1, category_l2 = get_closest_category(category_l1, category_l2)
    
    category = f"{category_l1}/{category_l2}"
    evidence = str(obj.get("evidence", "")).strip()
    conf_raw = obj.get("confidence", 0.0)
    try:
        confidence = float(conf_raw)
    except Exception:
        confidence = 0.0
    confidence = max(0.0, min(confidence, 1.0))
    return {
        "category_l1": category_l1,
        "category_l2": category_l2,
        "category": category,
        "confidence": confidence,
        "evidence": evidence,
    }


def _llm_classify_from_capsule_old(
    modality: str,
    summary: str,
    keyphrases: List[str],
    model: Optional[str],
    meta: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    """旧的分类逻辑（动态生成类别）"""
    system = (
        "你是一个企业资料库的文件分类助手。"
        "请基于提供的内容胶囊(capsule)，并结合文件元信息(如果提供)进行分类。"
        "输出必须是严格 JSON：{\"category_l1\": string, \"category_l2\": string, \"confidence\": number, \"evidence\": string}."
        "\n"
        "要求：\n"
        "- category_l1: 2-8 个中文字符，一级类目（例如：产品资料、生产采购、营销素材、安装售后、证书合规、其他）。\n"
        "- category_l2: 2-12 个中文字符，二级类目（例如：产品图片、规格页、BOM、入库表、安装说明、故障排查等）。\n"
        "- confidence: 0~1。\n"
        "- evidence: 引用 summary/keyphrases 或 文件元信息中的编号/关键字来说明理由。\n"
        "- 不要臆测：如果元信息与 summary/keyphrases 冲突或信息不足，请在 evidence 里说明不确定。"
    )

    meta_text = _format_meta_text(meta)
    user = (
        f"{meta_text}"
        f"模态: {modality}\n"
        f"summary: {summary}\n"
        f"keyphrases: {json.dumps(keyphrases, ensure_ascii=False)}\n"
        "请只返回 JSON。"
    )

    resp = _chat_json(
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        model=model,
    )

    content = ""
    try:
        content = resp["choices"][0]["message"]["content"]
    except Exception:
        content = str(resp)

    obj = _extract_json_obj(content)
    if obj is None:
        return {
            "category_l1": "未识别",
            "category_l2": "未识别",
            "category": "未识别/未识别",
            "confidence": 0.0,
            "evidence": (content or "")[:200],
        }

    category_l1 = str(obj.get("category_l1", "")).strip() or "未分类"
    category_l2 = str(obj.get("category_l2", "")).strip() or "未分类"
    category = f"{category_l1}/{category_l2}"
    evidence = str(obj.get("evidence", "")).strip()
    conf_raw = obj.get("confidence", 0.0)
    try:
        confidence = float(conf_raw)
    except Exception:
        confidence = 0.0
    confidence = max(0.0, min(confidence, 1.0))
    return {
        "category_l1": category_l1,
        "category_l2": category_l2,
        "category": category,
        "confidence": confidence,
        "evidence": evidence,
    }


def _load_existing_paths(jsonl_path: Path) -> Set[str]:
    if not jsonl_path.exists():
        return set()

    seen: Set[str] = set()
    try:
        with jsonl_path.open("r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                try:
                    obj = json.loads(line)
                except Exception:
                    continue
                p = str(obj.get("rel_path", "")).strip()
                if p:
                    seen.add(p)
    except Exception:
        return set()

    return seen


def _load_existing_file_ids(jsonl_path: Path) -> Set[str]:
    if not jsonl_path.exists():
        return set()

    seen: Set[str] = set()
    try:
        with jsonl_path.open("r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                try:
                    obj = json.loads(line)
                except Exception:
                    continue
                fid = str(obj.get("file_id", "")).strip()
                if fid:
                    seen.add(fid)
    except Exception:
        return set()
    return seen


def _cmd_build_capsules(args: argparse.Namespace) -> int:
    root = Path(args.root)
    out_path = Path(args.output)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    vision_model = (args.vision_model or "").strip() or None
    text_model = (args.text_model or "").strip() or None

    existing = _load_existing_file_ids(out_path) if args.resume else set()

    kinds_filter: Set[str] = set()
    if (args.kinds or "").strip():
        kinds_filter = {k.strip() for k in str(args.kinds).split(",") if k.strip()}

    items: List[FileItem] = list(_iter_files(root))
    items.sort(key=lambda x: x.rel_path)
    if kinds_filter:
        items = [it for it in items if _kind_for_ext(it.ext) in kinds_filter]
    if args.limit and args.limit > 0:
        items = items[: args.limit]

    now = datetime.utcnow().isoformat() + "Z"
    processed = 0
    skipped = 0
    progress_every = int(getattr(args, "progress_every", 0) or 0)
    t0 = time.time()
    total = len(items)

    with out_path.open("a", encoding="utf-8") as f:
        for idx, it in enumerate(items, start=1):
            file_id = it.rel_path
            if file_id in existing:
                skipped += 1
                continue

            kind = _kind_for_ext(it.ext)
            if kinds_filter and kind not in kinds_filter:
                continue
            file_name = it.abs_path.name

            meta = {
                "rel_path": it.rel_path,
                "file_name": file_name,
                "ext": it.ext,
                "size": it.size,
                "mtime": it.mtime,
            }

            modality = "document"
            capsule: Dict[str, Any] = {}

            if kind == "image":
                modality = "image"
                data_url = _to_data_url_for_image(it.abs_path, max_bytes=int(args.max_image_bytes))
                if data_url is None:
                    capsule = {"summary": "图片过大或读取失败，未生成视觉摘要。", "keyphrases": [], "confidence_read": 0.0}
                else:
                    capsule = _vision_capsule(data_url=data_url, model=vision_model, meta=meta)
            elif kind == "excel":
                modality = "table"
                excerpt = _read_excel_snippet(it.abs_path, max_chars=int(args.max_chars), nrows=int(args.excel_nrows))
                capsule = _text_capsule(kind=kind, excerpt=excerpt, model=text_model, meta=meta)
            elif kind == "text":
                excerpt = _read_text_snippet(it.abs_path, max_chars=int(args.max_chars))
                capsule = _text_capsule(kind=kind, excerpt=excerpt, model=text_model, meta=meta)
            elif kind == "pdf":
                excerpt, first_image = _read_pdf_excerpt_and_image(
                    it.abs_path,
                    max_chars=int(args.max_chars),
                    max_pages=int(args.pdf_max_pages),
                    lang=str(args.pdf_lang),
                    max_image_bytes=int(args.max_image_bytes),
                )
                if (excerpt or "").strip():
                    capsule = _text_capsule(kind=kind, excerpt=excerpt, model=text_model, meta=meta)
                elif (first_image or "").strip():
                    capsule = _vision_capsule(data_url=str(first_image), model=vision_model, meta=meta)
                else:
                    capsule = {"summary": "PDF内容抽取失败或为空，未生成内容摘要。", "keyphrases": [], "confidence_read": 0.0}
            elif kind == "docx":
                excerpt = _read_docx_snippet(it.abs_path, max_chars=int(args.max_chars))
                if (excerpt or "").strip():
                    capsule = _text_capsule(kind=kind, excerpt=excerpt, model=text_model, meta=meta)
                else:
                    img = _extract_first_docx_image_data_url(it.abs_path, max_bytes=int(args.max_image_bytes))
                    if (img or "").strip():
                        capsule = _vision_capsule(data_url=str(img), model=vision_model, meta=meta)
                    else:
                        capsule = {"summary": "DOCX内容抽取失败或为空，未生成内容摘要。", "keyphrases": [], "confidence_read": 0.0}
            else:
                capsule = {"summary": "暂不支持的文件类型，未生成内容摘要。", "keyphrases": [], "confidence_read": 0.0}

            rec = {
                "ts": now,
                "file_id": file_id,
                "rel_path": it.rel_path,
                "file_name": file_name,
                "ext": it.ext,
                "kind": kind,
                "modality": modality,
                "size": it.size,
                "mtime": it.mtime,
                **capsule,
            }

            f.write(json.dumps(rec, ensure_ascii=False) + "\n")
            f.flush()
            processed += 1

            if progress_every > 0 and (processed % progress_every == 0):
                elapsed = max(0.001, time.time() - t0)
                rate = processed / elapsed
                done = idx
                remaining = max(0, total - done)
                eta = remaining / rate if rate > 0 else 0.0
                eta_min = int(eta // 60)
                print(
                    f"progress build-capsules: processed={processed} skipped={skipped} "
                    f"idx={done}/{total} rate={rate:.2f}/s elapsed={elapsed:.1f}s eta~{eta_min}m "
                    f"file={it.rel_path}"
                )

    print(f"root={root}")
    print(f"output={out_path}")
    print(f"processed={processed} skipped={skipped}")
    return 0


def _cmd_diagnose_vision(args: argparse.Namespace) -> int:
    root = Path(args.root)
    vision_model = (args.vision_model or "").strip() or None

    items: List[FileItem] = []
    for it in _iter_files(root):
        if _kind_for_ext(it.ext) != "image":
            continue
        items.append(it)
    items.sort(key=lambda x: x.rel_path)

    n = int(args.count)
    if n <= 0:
        n = 3
    items = items[:n]

    if not items:
        print(f"No images found under: {root}")
        return 1

    for idx, it in enumerate(items, start=1):
        data_url = _to_data_url_for_image(it.abs_path, max_bytes=int(args.max_image_bytes))
        print("=" * 80)
        print(f"[{idx}/{len(items)}] rel_path={it.rel_path}")
        print(f"  size={it.size} bytes")
        requested = _normalize_dashscope_model_name(_resolve_vision_model(vision_model))
        print(f"  requested_model={requested}")

        if data_url is None:
            print("  data_url=None (image too large or read failed)")
            continue

        print(f"  data_url_len={len(data_url)}")

        system = (
            "你是一个图片内容分析助手。"
            "请描述图片里最主要的对象/场景/文字（如果有），不要臆测。"
            "输出可以是自然语言。"
        )
        # Keep ordering consistent with DashScope OpenAI-compatible examples.
        meta_text = _format_meta_text(
            {
                "rel_path": it.rel_path,
                "file_name": it.abs_path.name,
                "ext": it.ext,
                "size": it.size,
                "mtime": it.mtime,
            }
        )
        user = [
            {"type": "image_url", "image_url": {"url": data_url}},
            {"type": "text", "text": f"{meta_text}请直接描述图片内容。"},
        ]

        client = _get_openai_client()
        try:
            resp = client.chat.completions.create(
                model=requested,
                messages=[
                    {"role": "system", "content": system},
                    {"role": "user", "content": user},
                ],
            )
        except Exception as e:
            print(f"  completion_error={repr(e)}")
            continue

        print(f"  response_id={getattr(resp, 'id', None)}")
        print(f"  response_model={getattr(resp, 'model', None)}")
        print(f"  usage={getattr(resp, 'usage', None)}")

        content = (resp.choices[0].message.content or "").strip()

        snippet = (content or "").strip().replace("\n", " ")
        if len(snippet) > 500:
            snippet = snippet[:500]
        print(f"  content_snippet={snippet}")

        if _response_suggests_missing_image(content):
            fb = _resolve_fallback_vision_model(_resolve_vision_model(vision_model))
            print(f"  warning=vision_model_may_ignore_image_payload fallback_suggestion={fb}")

    return 0


def _cmd_classify(args: argparse.Namespace) -> int:
    in_path = Path(args.input)
    out_path = Path(args.output)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    model = (args.model or "").strip() or None
    existing = _load_existing_file_ids(out_path) if args.resume else set()

    processed = 0
    skipped = 0
    counts: Dict[str, int] = {}
    counts_l1: Dict[str, int] = {}

    progress_every = int(getattr(args, "progress_every", 0) or 0)
    t0 = time.time()

    now = datetime.utcnow().isoformat() + "Z"

    with in_path.open("r", encoding="utf-8") as fin, out_path.open("a", encoding="utf-8") as fout:
        for idx, line in enumerate(fin, start=1):
            line = line.strip()
            if not line:
                continue
            try:
                cap = json.loads(line)
            except Exception:
                continue

            file_id = str(cap.get("file_id", "")).strip()
            if not file_id:
                continue
            if file_id in existing:
                skipped += 1
                continue

            modality = str(cap.get("modality", "")).strip() or "document"
            summary = str(cap.get("summary", "")).strip()
            keyphrases_raw = cap.get("keyphrases", [])
            keyphrases: List[str] = []
            if isinstance(keyphrases_raw, list):
                keyphrases = [str(x).strip() for x in keyphrases_raw if str(x).strip()]

            meta = {
                "rel_path": cap.get("rel_path"),
                "file_name": cap.get("file_name"),
                "ext": cap.get("ext"),
                "size": cap.get("size"),
                "mtime": cap.get("mtime"),
            }

            cls = _llm_classify_from_capsule(
                modality=modality,
                summary=summary,
                keyphrases=keyphrases,
                model=model,
                meta=meta,
            )

            rec = {"ts": now, **cap, **cls}
            fout.write(json.dumps(rec, ensure_ascii=False) + "\n")
            fout.flush()
            processed += 1
            c = str(rec.get("category", ""))
            counts[c] = counts.get(c, 0) + 1
            c1 = str(rec.get("category_l1", ""))
            counts_l1[c1] = counts_l1.get(c1, 0) + 1

            if progress_every > 0 and (processed % progress_every == 0):
                elapsed = max(0.001, time.time() - t0)
                rate = processed / elapsed
                eta_min = 0
                if rate > 0:
                    # We don't know the total line count cheaply; show processed + rate.
                    eta_min = 0
                print(
                    f"progress classify: processed={processed} skipped={skipped} "
                    f"idx={idx} rate={rate:.2f}/s elapsed={elapsed:.1f}s eta~=unknown "
                    f"file={file_id}"
                )

    print(f"input={in_path}")
    print(f"output={out_path}")
    print(f"processed={processed} skipped={skipped}")
    print("categories_l1:")
    for k, v in sorted(counts_l1.items(), key=lambda kv: (-kv[1], kv[0])):
        print(f"  {k}: {v}")
    print("categories_l1_l2:")
    for k, v in sorted(counts.items(), key=lambda kv: (-kv[1], kv[0])):
        print(f"  {k}: {v}")
    return 0


def _cmd_run_all(args: argparse.Namespace) -> int:
    caps_args = argparse.Namespace(
        root=args.root,
        output=args.capsules_output,
        limit=int(args.limit),
        kinds=str(args.kinds or ""),
        text_model=str(args.text_model or ""),
        vision_model=str(args.vision_model or ""),
        progress_every=int(args.progress_every),
        max_chars=int(args.max_chars),
        excel_nrows=int(args.excel_nrows),
        max_image_bytes=int(args.max_image_bytes),
        pdf_max_pages=int(args.pdf_max_pages),
        pdf_lang=str(args.pdf_lang),
        resume=bool(args.resume),
    )
    rc = _cmd_build_capsules(caps_args)
    if rc != 0:
        return rc

    cls_args = argparse.Namespace(
        input=args.capsules_output,
        output=args.classified_output,
        limit=int(args.classify_limit),
        model=str(args.classify_model or ""),
        progress_every=int(args.progress_every),
        resume=bool(args.resume),
    )
    return _cmd_classify(cls_args)


def main(argv: Optional[Sequence[str]] = None) -> int:
    ap = argparse.ArgumentParser()
    sub = ap.add_subparsers(dest="cmd", required=True)

    ap_caps = sub.add_parser("build-capsules")
    ap_caps.add_argument("--root", default=str(Path(__file__).resolve().parents[1] / "structured_processed_files" / "🔥Spa"))
    ap_caps.add_argument("--output", default=str(Path(__file__).resolve().parents[1] / "data_storage" / "spa_capsules.jsonl"))
    ap_caps.add_argument("--limit", type=int, default=20)
    ap_caps.add_argument("--kinds", default="")
    ap_caps.add_argument("--text-model", default="")
    ap_caps.add_argument("--vision-model", default="")
    ap_caps.add_argument("--max-chars", type=int, default=8000)
    ap_caps.add_argument("--excel-nrows", type=int, default=20)
    ap_caps.add_argument("--max-image-bytes", type=int, default=2_500_000)
    ap_caps.add_argument("--pdf-max-pages", type=int, default=3)
    ap_caps.add_argument("--pdf-lang", default="en")
    ap_caps.add_argument("--resume", action="store_true")
    ap_caps.add_argument("--progress-every", type=int, default=50)

    ap_cls = sub.add_parser("classify")
    ap_cls.add_argument("--input", default=str(Path(__file__).resolve().parents[1] / "data_storage" / "spa_capsules.jsonl"))
    ap_cls.add_argument("--output", default=str(Path(__file__).resolve().parents[1] / "data_storage" / "spa_classified.jsonl"))
    ap_cls.add_argument("--limit", type=int, default=0)
    ap_cls.add_argument("--model", default="")
    ap_cls.add_argument("--resume", action="store_true")
    ap_cls.add_argument("--progress-every", type=int, default=50)

    ap_all = sub.add_parser("run-all")
    ap_all.add_argument("--root", default=str(Path(__file__).resolve().parents[1] / "structured_processed_files" / "🔥Spa"))
    ap_all.add_argument("--capsules-output", default=str(Path(__file__).resolve().parents[1] / "data_storage" / "spa_capsules_all.jsonl"))
    ap_all.add_argument("--classified-output", default=str(Path(__file__).resolve().parents[1] / "data_storage" / "spa_classified_hier_all.jsonl"))
    ap_all.add_argument("--limit", type=int, default=0)
    ap_all.add_argument("--kinds", default="")
    ap_all.add_argument("--text-model", default="")
    ap_all.add_argument("--vision-model", default="")
    ap_all.add_argument("--classify-model", default="")
    ap_all.add_argument("--classify-limit", type=int, default=0)
    ap_all.add_argument("--progress-every", type=int, default=50)
    ap_all.add_argument("--max-chars", type=int, default=8000)
    ap_all.add_argument("--excel-nrows", type=int, default=20)
    ap_all.add_argument("--max-image-bytes", type=int, default=6_000_000)
    ap_all.add_argument("--pdf-max-pages", type=int, default=3)
    ap_all.add_argument("--pdf-lang", default="en")
    ap_all.add_argument("--resume", action="store_true")

    ap_diag = sub.add_parser("diagnose-vision")
    ap_diag.add_argument("--root", default=str(Path(__file__).resolve().parents[1] / "structured_processed_files" / "🔥Spa"))
    ap_diag.add_argument("--vision-model", default="")
    ap_diag.add_argument("--count", type=int, default=3)
    ap_diag.add_argument("--max-image-bytes", type=int, default=6_000_000)

    args = ap.parse_args(list(argv) if argv is not None else None)

    if load_dotenv is not None:
        try:
            load_dotenv(Path(__file__).resolve().parents[1] / ".env")
        except Exception:
            pass

    if args.cmd == "build-capsules":
        return _cmd_build_capsules(args)
    if args.cmd == "classify":
        return _cmd_classify(args)
    if args.cmd == "run-all":
        return _cmd_run_all(args)
    if args.cmd == "diagnose-vision":
        return _cmd_diagnose_vision(args)
    raise SystemExit(2)


if __name__ == "__main__":
    raise SystemExit(main())
