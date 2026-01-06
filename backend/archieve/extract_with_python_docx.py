#!/usr/bin/env python3
"""
仅使用python-docx库提取test.docx文件内容
"""

import os
import json
from docx import Document
import zipfile


def extract_text_with_python_docx(file_path: str):
    """
    使用python-docx提取docx文件的文本内容，按文档顺序排列
    
    Args:
        file_path (str): docx文件路径
    
    Returns:
        dict: 包含文档内容的字典
    """
    if not os.path.exists(file_path):
        print(f"错误: 文件 {file_path} 不存在")
        return None
    
    try:
        # 打开docx文件
        doc = Document(file_path)
        
        print(f"正在提取文件: {file_path}")
        
        # 提取文档信息
        doc_info = {
            "filename": os.path.basename(file_path),
            "paragraphs_count": len(doc.paragraphs),
            "tables_count": len(doc.tables),
            "sections_count": len(doc.sections)
        }
        
        # 按顺序提取所有内容元素
        ordered_content = []
        
        # 遍历文档的所有元素（段落和表格）
        for element in doc.element.body:
            if element.tag.endswith('p'):  # 段落
                # 找到对应的段落对象
                for para in doc.paragraphs:
                    if para._element == element:
                        if para.text.strip():  # 只保存非空段落
                            para_info = {
                                "type": "paragraph",
                                "index": len(ordered_content) + 1,
                                "text": para.text.strip(),
                                "style": para.style.name if para.style else "Normal",
                                "alignment": str(para.alignment) if para.alignment else "None"
                            }
                            ordered_content.append(para_info)
                        break
            
            elif element.tag.endswith('tbl'):  # 表格
                # 找到对应的表格对象
                for table in doc.tables:
                    if table._element == element:
                        table_data = {
                            "type": "table",
                            "index": len(ordered_content) + 1,
                            "rows": len(table.rows),
                            "columns": len(table.columns),
                            "cells": []
                        }
                        
                        for row_idx, row in enumerate(table.rows):
                            row_data = []
                            for cell_idx, cell in enumerate(row.cells):
                                cell_info = {
                                    "row": row_idx + 1,
                                    "column": cell_idx + 1,
                                    "text": cell.text.strip()
                                }
                                row_data.append(cell_info)
                            table_data["cells"].append(row_data)
                        
                        ordered_content.append(table_data)
                        break
        
        # 提取图片信息（用于参考）
        images = extract_images_from_docx(file_path)
        
        # 组合所有内容
        content = {
            "document_info": doc_info,
            "ordered_content": ordered_content,
            "images": images
        }
        
        print(f"提取完成:")
        print(f"  总元素数: {len(ordered_content)}")
        print(f"  图片数: {len(images)}")
        
        return content
        
    except Exception as e:
        print(f"提取文件时发生错误: {e}")
        return None


def extract_images_from_docx(file_path: str):
    """
    从docx文件中提取图片信息
    
    Args:
        file_path (str): docx文件路径
    
    Returns:
        list: 图片信息列表
    """
    images = []
    
    try:
        # 打开docx文件作为zip文件
        with zipfile.ZipFile(file_path, 'r') as docx_zip:
            # 查找图片文件
            image_files = [f for f in docx_zip.namelist() if f.startswith('word/media/') and f != 'word/media/']
            
            for img_file in image_files:
                try:
                    # 读取图片数据
                    img_data = docx_zip.read(img_file)
                    
                    # 获取文件扩展名
                    _, ext = os.path.splitext(img_file)
                    if not ext:
                        ext = '.jpeg'  # 默认扩展名
                    
                    # 创建图片信息字典
                    image_info = {
                        "filename": os.path.basename(img_file),
                        "path": img_file,
                        "size": len(img_data),
                        "extension": ext,
                        "size_mb": round(len(img_data) / (1024 * 1024), 2)
                    }
                    images.append(image_info)
                    
                except Exception as e:
                    print(f"读取图片 {img_file} 时出错: {e}")
        
        return images
        
    except Exception as e:
        print(f"提取图片时出错: {e}")
        return []


def save_images_to_directory(images, output_dir: str = "python_docx_images"):
    """
    保存图片到指定目录
    
    Args:
        images: 图片信息列表
        output_dir (str): 输出目录路径
    """
    if not images:
        print("没有图片可保存")
        return
    
    try:
        # 创建输出目录
        os.makedirs(output_dir, exist_ok=True)
        
        # 重新打开docx文件提取图片数据
        with zipfile.ZipFile("test.docx", 'r') as docx_zip:
            saved_count = 0
            for i, img in enumerate(images, 1):
                try:
                    # 生成新的文件名
                    new_filename = f"docx_image_{i:03d}{img['extension']}"
                    new_path = os.path.join(output_dir, new_filename)
                    
                    # 读取并保存图片文件
                    img_data = docx_zip.read(img['path'])
                    with open(new_path, 'wb') as f:
                        f.write(img_data)
                    
                    print(f"图片已保存: {new_path} ({img['size_mb']} MB)")
                    saved_count += 1
                    
                except Exception as e:
                    print(f"保存图片 {img['filename']} 时出错: {e}")
        
        print(f"成功保存 {saved_count} 张图片到目录: {output_dir}")
        
    except Exception as e:
        print(f"保存图片时发生错误: {e}")


def display_content(content):
    """
    显示提取的内容，按文档顺序排列
    
    Args:
        content: 提取的内容字典
    """
    if not content:
        print("没有内容可显示")
        return
    
    print("\n" + "="*60)
    print("文档信息:")
    print("="*60)
    doc_info = content["document_info"]
    print(f"文件名: {doc_info['filename']}")
    print(f"段落数: {doc_info['paragraphs_count']}")
    print(f"表格数: {doc_info['tables_count']}")
    print(f"节数: {doc_info['sections_count']}")
    
    print("\n" + "="*60)
    print("按文档顺序的内容:")
    print("="*60)
    
    # 统计各类型元素数量
    paragraph_count = 0
    table_count = 0
    
    for i, element in enumerate(content["ordered_content"][:20], 1):  # 只显示前20个元素
        if element["type"] == "paragraph":
            paragraph_count += 1
            print(f"\n[{i}] 段落 {paragraph_count}:")
            print(f"  样式: {element['style']}")
            print(f"  对齐: {element['alignment']}")
            print(f"  内容: {element['text'][:100]}{'...' if len(element['text']) > 100 else ''}")
        
        elif element["type"] == "table":
            table_count += 1
            print(f"\n[{i}] 表格 {table_count}:")
            print(f"  行数: {element['rows']}, 列数: {element['columns']}")
            # 显示表格前几行
            for row_idx, row in enumerate(element["cells"][:2]):
                row_text = " | ".join([cell["text"] for cell in row])
                print(f"  行 {row_idx + 1}: {row_text[:80]}{'...' if len(row_text) > 80 else ''}")
            
            if element["rows"] > 2:
                print(f"  ... 还有 {element['rows'] - 2} 行")
    
    if len(content["ordered_content"]) > 20:
        print(f"\n... 还有 {len(content['ordered_content']) - 20} 个元素")
    
    print("\n" + "="*60)
    print("图片信息:")
    print("="*60)
    for img in content["images"]:
        print(f"  图片: {img['filename']} ({img['size_mb']} MB)")
    
    # 显示元素类型统计
    print("\n" + "="*60)
    print("元素类型统计:")
    print("="*60)
    total_elements = len(content["ordered_content"])
    paragraphs_in_order = sum(1 for elem in content["ordered_content"] if elem["type"] == "paragraph")
    tables_in_order = sum(1 for elem in content["ordered_content"] if elem["type"] == "table")
    
    print(f"总元素数: {total_elements}")
    print(f"段落数: {paragraphs_in_order}")
    print(f"表格数: {tables_in_order}")
    print(f"图片数: {len(content['images'])}")


def save_to_json(content, output_file: str):
    """
    将内容保存为JSON格式
    
    Args:
        content: 提取的内容字典
        output_file (str): 输出文件路径
    """
    if not content:
        print("没有内容可保存")
        return
    
    try:
        # 保存到文件，确保中文正确显示
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(content, f, ensure_ascii=False, indent=2)
        
        print(f"内容已保存到: {output_file}")
        
    except Exception as e:
        print(f"保存JSON文件时发生错误: {e}")


def main():
    """主函数"""
    # 文件路径
    docx_file = "test.docx"
    json_output = "python_docx_content.json"
    images_output_dir = "python_docx_images"
    
    print("开始使用python-docx提取test.docx文件内容...")
    
    # 提取内容
    content = extract_text_with_python_docx(docx_file)
    
    if content:
        # 显示内容
        display_content(content)
        
        # 保存图片
        if content["images"]:
            save_images_to_directory(content["images"], images_output_dir)
        
        # 保存为JSON
        save_to_json(content, json_output)
        
        print(f"\n提取完成!")
        total_elements = len(content['ordered_content'])
        paragraphs_count = sum(1 for elem in content['ordered_content'] if elem['type'] == 'paragraph')
        tables_count = sum(1 for elem in content['ordered_content'] if elem['type'] == 'table')
        print(f"总元素数: {total_elements}")
        print(f"文本段落: {paragraphs_count} 个")
        print(f"表格: {tables_count} 个")
        print(f"图片: {len(content['images'])} 张")
    else:
        print("提取失败")


if __name__ == "__main__":
    main()
